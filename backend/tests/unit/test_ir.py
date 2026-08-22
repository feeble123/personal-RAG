"""P1-1 单元2：DocumentElement IR 重构。

覆盖：
- ir_validation：非法 bbox / 乱序 reading_order / parent 悬空 / 表格行列不一致 → 拒绝
- ParsedBlock.to_element：text/page/section 不丢、heading 标记 inferred
- 各 parser 产出 elements：100% 过 validator、section_path 非空（有 section 的块）
- parse_to_elements：有 elements 用 elements，无则回退 blocks→elements
"""
from __future__ import annotations

from pathlib import Path

import pytest

from app.services.parser.base import ParsedBlock
from app.services.parser.factory import get_parser, parse_to_elements
from app.services.parser.ir import DocumentElement, ElementType
from app.services.parser.ir_validation import validate_elements


class TestValidator:
    def test_valid_elements_pass(self):
        els = [
            DocumentElement(element_id="a-0", type=ElementType.PARAGRAPH, text="正文", reading_order=0),
            DocumentElement(element_id="a-1", type=ElementType.PARAGRAPH, text="正文2", reading_order=1),
        ]
        assert validate_elements(els) == []

    def test_duplicate_id(self):
        els = [
            DocumentElement(element_id="a-0", type=ElementType.PARAGRAPH, text="x", reading_order=0),
            DocumentElement(element_id="a-0", type=ElementType.PARAGRAPH, text="y", reading_order=1),
        ]
        errs = validate_elements(els)
        assert any("重复" in e for e in errs)

    def test_bad_bbox(self):
        el = DocumentElement(
            element_id="a-0", type=ElementType.PARAGRAPH, text="x",
            reading_order=0, bbox=(10, 10, 5, 5),  # x1 < x0
        )
        errs = validate_elements([el])
        assert any("bbox" in e for e in errs)

    def test_reading_order_mismatch(self):
        el = DocumentElement(
            element_id="a-0", type=ElementType.PARAGRAPH, text="x", reading_order=5,
        )
        errs = validate_elements([el])
        assert any("reading_order" in e for e in errs)

    def test_hanging_parent(self):
        el = DocumentElement(
            element_id="a-0", type=ElementType.PARAGRAPH, text="x",
            reading_order=0, parent_id="does-not-exist",
        )
        errs = validate_elements([el])
        assert any("parent_id" in e for e in errs)

    def test_bad_heading_level(self):
        el = DocumentElement(
            element_id="a-0", type=ElementType.HEADING, text="标题",
            reading_order=0, heading_level=9,
        )
        errs = validate_elements([el])
        assert any("heading_level" in e for e in errs)

    def test_table_row_mismatch(self):
        el = DocumentElement(
            element_id="a-0", type=ElementType.TABLE, text="表",
            reading_order=0,
            table={"rows": [["a", "b", "c"], ["d"]], "header_path": ["h1", "h2"]},
        )
        errs = validate_elements([el])
        assert any("table" in e for e in errs)

    def test_page_range_invalid(self):
        el = DocumentElement(
            element_id="a-0", type=ElementType.PARAGRAPH, text="x",
            reading_order=0, page_start=3, page_end=1,
        )
        errs = validate_elements([el])
        assert any("page_start" in e for e in errs)

    def test_empty_text_rejected(self):
        el = DocumentElement(
            element_id="a-0", type=ElementType.PARAGRAPH, text="  ", reading_order=0,
        )
        errs = validate_elements([el])
        assert any("text" in e for e in errs)


class TestBlockToElement:
    def test_paragraph_roundtrip(self):
        b = ParsedBlock(text="正文内容", section="第一章 / 1.1 节", page=3)
        el = b.to_element(0, "md")
        assert el.type == ElementType.PARAGRAPH
        assert el.text == "正文内容"
        assert el.section_path == ("第一章", "1.1 节")
        assert el.page_start == 3
        assert el.reading_order == 0
        assert el.source_ref["parser"] == "md"
        assert el.source_ref["block_index"] == 0

    def test_heading_marks_inferred(self):
        b = ParsedBlock(text="1 总则", block_type="heading", page=1)
        el = b.to_element(0, "pdf")
        assert el.type == ElementType.HEADING
        assert "inferred_heading" in el.flags

    def test_heading_with_level_no_inferred(self):
        b = ParsedBlock(text="1 总则", block_type="heading", page=1)
        el = b.to_element(0, "pdf", heading_level=1)
        assert el.heading_level == 1
        assert "inferred_heading" not in el.flags

    def test_table_type(self):
        b = ParsedBlock(text="a | b", section="表", block_type="table")
        el = b.to_element(0, "excel")
        assert el.type == ElementType.TABLE


class TestParsersProduceValidIR:
    def test_markdown_elements_valid(self, tmp_path):
        p = tmp_path / "t.md"
        p.write_text("# 第一章 总则\n\n正文段落。\n\n## 1.1 节\n\n第二节内容。\n", encoding="utf-8")
        parsed = get_parser("t.md").parse(p, "t.md")
        assert parsed.elements, "md 应产出 elements"
        errs = validate_elements(parsed.elements)
        assert not errs, f"md IR 应全合法: {errs}"
        # 有标题块 + 段落块
        types = {el.type for el in parsed.elements}
        assert ElementType.HEADING in types
        assert ElementType.PARAGRAPH in types

    def test_text_elements_valid(self, tmp_path):
        p = tmp_path / "t.txt"
        p.write_text("第一段。\n\n第二段。\n", encoding="utf-8")
        parsed = get_parser("t.txt").parse(p, "t.txt")
        errs = validate_elements(parsed.elements)
        assert not errs, f"txt IR 应全合法: {errs}"

    def test_csv_elements_valid(self, tmp_path):
        p = tmp_path / "t.csv"
        p.write_text("名称,数量\n甲,1\n乙,2\n", encoding="utf-8-sig")
        parsed = get_parser("t.csv").parse(p, "t.csv")
        errs = validate_elements(parsed.elements)
        assert not errs, f"csv IR 应全合法: {errs}"
        assert all(el.type == ElementType.TABLE for el in parsed.elements)

    def test_docx_elements_valid(self, tmp_path):
        """构造一个最小 docx（python-docx）。"""
        pytest.importorskip("docx")
        from docx import Document

        p = tmp_path / "t.docx"
        doc = Document()
        doc.add_heading("第一章 总则", level=1)
        doc.add_paragraph("正文段落。")
        doc.add_heading("1.1 节", level=2)
        doc.add_paragraph("第二节内容。")
        doc.save(str(p))
        parsed = get_parser("t.docx").parse(p, "t.docx")
        errs = validate_elements(parsed.elements)
        assert not errs, f"docx IR 应全合法: {errs}"
        types = {el.type for el in parsed.elements}
        assert ElementType.HEADING in types

    def test_parse_to_elements_returns_valid(self, tmp_path):
        p = tmp_path / "t.md"
        p.write_text("# 总则\n\n正文。\n", encoding="utf-8")
        els = parse_to_elements(p, "t.md")
        errs = validate_elements(els)
        assert not errs, f"parse_to_elements IR 应全合法: {errs}"
