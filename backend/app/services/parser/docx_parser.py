"""Word .docx 解析：段落 + Heading 样式树 + 表格，按文档元素顺序输出。"""
from __future__ import annotations

import logging
from pathlib import Path

from docx import Document  # type: ignore[import-not-found]

from app.services.parser.base import DocumentParser, ParsedBlock, ParsedDocument

logger = logging.getLogger(__name__)


def _paragraph_heading_level(text: str, style_name: str | None) -> int | None:
    """返回标题级别（1-9）或 None（正文）。

    P1-4 最小集：支持中文本地化样式名（「标题 1」「标题一」），不只英文 Heading N。
    """
    if not style_name:
        return None
    s = style_name.strip()
    low = s.lower()
    # 英文 Heading N / HeadingN
    if low.startswith("heading"):
        digits = low.replace("heading", "").strip()
        if digits.isdigit():
            return int(digits)
        return 1
    # 中文「标题 1」/「标题一」/「标题1」
    if s.startswith("标题"):
        rest = s.replace("标题", "").strip()
        if rest.isdigit():
            return int(rest)
        cn = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6}
        if rest in cn:
            return cn[rest]
    return None


def _docx_heading_level(block: ParsedBlock) -> int | None:
    """docx 标题块层级（IR 用）：blocks 未携带原始 style 层级，用编号模式估。"""
    if block.block_type != "heading":
        return None
    from app.services.parser.headings import heading_level

    lvl = heading_level(block.text)
    return lvl if lvl >= 1 else None


class DocxParser(DocumentParser):
    extensions = ("docx",)

    def parse(self, path: Path, filename: str, chunk_strategy: str = "old") -> ParsedDocument:
        doc = Document(str(path))
        blocks: list[ParsedBlock] = []
        section_stack: list[str] = []
        quality: dict = {"parser": "docx", "tables": 0, "headings": 0, "paragraphs": 0}

        # 标题映射：Heading N → 章节路径
        from docx.oxml.ns import qn

        p_map = {p._p: p for p in doc.paragraphs}
        t_map = {t._tbl: t for t in doc.tables}

        for child in doc.element.body.iterchildren():
            if child.tag == qn("w:p"):
                para = p_map.get(child)
                if para is None or not para.text.strip():
                    continue
                text = para.text.strip()
                level = _paragraph_heading_level(text, para.style.name if para.style else None)
                if level:
                    # 章节路径：覆盖/追加
                    if len(section_stack) >= level:
                        section_stack = section_stack[: level - 1]
                    section_stack.append(text)
                    quality["headings"] += 1
                    blocks.append(
                        ParsedBlock(
                            text=text,
                            section="/".join(section_stack[:-1]) or None,
                            block_type="heading",
                        )
                    )
                else:
                    quality["paragraphs"] += 1
                    blocks.append(
                        ParsedBlock(
                            text=text,
                            section="/".join(section_stack) or None,
                            block_type="paragraph",
                        )
                    )
            elif child.tag == qn("w:tbl"):
                table = t_map.get(child)
                if table is None:
                    continue
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                rows = [r for r in rows if r.strip(" |")]
                if rows:
                    quality["tables"] += 1
                    blocks.append(
                        ParsedBlock(
                            text="\n".join(rows),
                            section="/".join(section_stack) or None,
                            block_type="table",
                        )
                    )

        quality["blocks"] = len(blocks)
        # P1-1：blocks → IR elements（docx 无页面概念，page/bbox 为 None）
        elements = [
            b.to_element(i, "docx", heading_level=_docx_heading_level(b))
            for i, b in enumerate(blocks)
        ]
        return ParsedDocument(blocks=blocks, quality=quality, elements=elements)
